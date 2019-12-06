# -*- encoding: utf-8 -*-
import os
import re
import copy
import chardet
import traceback
import codecs
import shutil
from docx import Document
from win32com.client import Dispatch
import openpyxl
from openpyxl import load_workbook
import random
from utils import Util
import docx
# import sys
# reload(sys)
# sys.setdefaultencoding('utf-8')

class FileSplitAndSplicing(object):
    """
    将txt或docx的文章拆分下来首段,中段,尾段,然后其按照之前的逻辑拼接文章.
    """

    def __init__(self, path):
        self.util = Util()
        self.document = docx
        # 读取
        self.read_path = path
        self.domain_name = "_".join(path.split("/")[-1].split("_")[:3])

        self.start_keyword = 0  # 关键词开始的位置
        self.end_keyword = int(path.split("/")[-1].split("_")[-1])  # 关键词结束的位置
        # self.special_keyword = '苏州'
        self.used_keyword = []

        # 保存路径
        self.save_article_path = r'./data/save_path/{}_articles'.format(self.domain_name)
        self.save_img_path = r'./data/save_path/{}_imgs'.format(self.domain_name)
        ####################   打包   ##########################
        # self.save_article_path = r'../data/save_path/{}_articles'.format(self.domain_name)
        # self.save_img_path = r'../data/save_path/{}_imgs'.format(self.domain_name)
        ########################################################

    # def get_keywords(self):
    #     """
    #     获取关键词
    #     举例:现需要150个关键词,含有"苏州"的关键词优先,
    #     如果含有"苏州"的关键词超过150个,则取前150个,
    #     如果含有"苏州"的关键词不到150,则取完这些词还要再取一些普通关键词凑够150个
    #     """
    #     all_keywords = self.read_xlsx()  # 所有的关键词
    #     special_keywords = [kw for kw in all_keywords if self.special_keyword in kw]  # 特殊关键词,如:含有"苏州"的关键词
    #     all_keywords = [kw for kw in all_keywords if self.special_keyword not in kw]  # 不含有"苏州"的所有的关键词
    #
    #     keywords_num = self.end_keyword - self.start_keyword  # 需要的关键词个数
    #     if len(special_keywords) >= keywords_num:
    #         needed_keywords = special_keywords[0:keywords_num]
    #         remaining_keywords = special_keywords[keywords_num:]
    #         remaining_keywords.extend(all_keywords)
    #     else:
    #         needed_keywords = copy.deepcopy(special_keywords)
    #         needed_keywords.extend(all_keywords[0: keywords_num-len(special_keywords)])
    #         remaining_keywords = all_keywords[(keywords_num-len(special_keywords)):]
    #     return needed_keywords, remaining_keywords

    # def get_file_list(self, filepath):
    #     '''获取（"反射型光电传感器"）目录下的所有文件及文件夹'''
    #     # return os.listdir(self.util.to_gbk(filepath))
    #     return os.listdir(filepath)

    def split_article(self, file_path, file):
        """
        将文章拆分为首段/中段/尾段
        """
        paragraphs = list()
        final_paragraphs = list()
        try:
            if file_path.endswith('txt'):
                try:
                    with open(file_path, 'r') as f:
                        paragraphs = f.readlines()
                except:
                    with open(file_path, 'r', encoding='UTF-8') as f:
                        paragraphs = f.readlines()
            elif file_path.endswith("docx"):
                document = Document(u'{}'.format(file_path))
                paragraphs = [p.text for p in document.paragraphs if p.text != '\n' and p != '' and p is not None]
            else:
                word = Dispatch('Word.Application')  # 打开word应用程序
                # word = DispatchEx('Word.Application') # 启动独立的进程
                word.Visible = 0  # 后台运行,不显示
                word.DisplayAlerts = 0  # 不警告
                dir_path = os.path.dirname(os.path.abspath(file)) + "\\" + file_path
                doc = word.Documents.Open(dir_path)
                # doc.SaveAs(os.path.splitext(dir_path)[0] + '.docx', 12, False, "", True, "", False, False, False, False)
                for para in doc.paragraphs:
                    paragraphs.append(para.Range.Text)
                doc.Close()

            # 处理文档中的非正常回车
            for p in paragraphs:
                ps = p.split("\n\t")
                for pp in ps:
                    final_paragraphs.append(pp)

            start = final_paragraphs[0]
            middle = final_paragraphs[1: -1]
            end = final_paragraphs[-1]
            return start, middle, end
        except:
            traceback.print_exc()

    def operate_picture(self, filepath):
        """
        处理图片
        :param filepath:
        :return: 所有图片的路径
        """
        imgs = []
        try:
            # for file in os.listdir(self.util.to_gbk(filepath)):
            #     img = file.decode('gbk').encode('utf-8')
            #     # imgs.append(os.path.join(filepath, img))
            #     imgs.append(img)
            for file in os.listdir(filepath):
                imgs.append(file)
            return imgs
        except:
            print('operate picture error', traceback.print_exc())

    def run(self):
        if not os.path.exists(self.save_article_path):
            os.mkdir(self.save_article_path)
        if not os.path.exists(self.save_img_path):
            os.mkdir(self.save_img_path)

        img_list = list()
        article_list = list()
        start_paragraph_list = list()  # 存放所有首段段落
        middle_paragraph_list = list()  # 存放所有中段段落
        end_paragraph_list = list()  # 存放所有尾段段落

        file_dir_list = self.util.get_file_dir(self.read_path)  # 获取所有文件

        for file in file_dir_list:
            file_path = self.read_path + "/" + file
            if file != 'img' and 'xlsx' not in file:
                start_paragraph, middle_paragraph, end_paragraph = self.split_article(file_path, file)
                start_paragraph_list.append(start_paragraph)
                middle_paragraph_list.extend(middle_paragraph)
                end_paragraph_list.append(end_paragraph)
            elif file == "img":
                img_list = self.operate_picture(file_path)  # 获取所有图片
                for img in img_list:
                    shutil.copy(u"{}".format(self.read_path + '/' + '/img/' + img), self.save_img_path)
            elif "xlsx" in file:
                self.keywords = self.util.read_xlsx(file_path)[self.start_keyword:self.end_keyword]  # 普通的取关键词

        # middle_paragraph_list = middle_paragraph_list if len(middle_paragraph_list) < 100 else random.sample(middle_paragraph_list, 100)
        all_mid_list = self.util.mid_permutation_and_combination(middle_paragraph_list)  # 中段所有排列组合之后的情况
        all_mid_list = all_mid_list if len(all_mid_list) < 2000 else random.sample(all_mid_list, 2000)
        articles = self.util.article_permutation_and_combination(start_paragraph_list, all_mid_list, end_paragraph_list)
        # articles = self.util.article_permutation_and_combination(random.sample(start_paragraph_list, 10), random.sample(all_mid_list, 10) , random.sample(end_paragraph_list, 10))
        article_list = self.util.get_article_list(articles, article_list)  # 存储最终的所有的文章【单个文件夹下的】

        # 下面每次循环生成一篇文章, 每个文件夹需要生成“every_article_num”篇文章
        # for _ in range(every_article_num):
        index = 1
        while True:
            keyword = self.util.get_keyword(self.keywords, self.used_keyword)
            if not keyword:  # 关键词使用完之后退出循环
                break
            print(index, keyword)

            # 随机抽取文章，要求文章字数在730~870
            # while True:
            #     article = random.choice(article_list)  # 随机抽一篇文章
            #     article_len = self.file_paocessing1.get_article_len(article)
            #     if 730 < article_len < 900:
            #         break

            try:
                article = random.choice(article_list)  # 随机抽一篇文章
                temp_article = copy.deepcopy(article)  # 深拷贝，对新数据进行处理，不改变原数据
                img = random.sample(img_list, 2)  # 随机取两张图
                article_str = ''
                #  段落 -- 对每一段进行处理
                for num in range(len(temp_article)):
                    if num == 0 or num == len(temp_article) - 1:  # 添加首段/尾段
                        temp_article[num] = self.util.insert_keyword(keyword, temp_article[num])  # 插入关键词
                        article_str += '<p>%s</p>\n' % temp_article[num]
                    elif num == 1:  # 添加第二段，并插入一张图片
                        article_str += '<p>%s</p>\n' % temp_article[num]
                        article_str += '<p><img src="{imgpath}/%s_imgs/%s"></p>\n' % (
                            self.domain_name, img[0])  # 注意修改站点名称
                    elif num == 3:  # 添加第四段，并插入一张图片
                        article_str += '<p>%s</p>\n' % temp_article[num]
                        article_str += '<p><img src="{imgpath}/%s_imgs/%s"></p>\n' % (
                            self.domain_name, img[1])  # 注意修改站点名称
                    else:  # 添加第三段
                        article_str += '<p>%s</p>\n' % temp_article[num]
                save_path = self.save_article_path + '/' + '{}.txt'.format(keyword)
                try:
                    self.util.write_article(save_path, article_str)
                except:
                    self.util.write_article(save_path, article_str.replace(u'\u200b', u'').replace(u'\xa0', u'').replace(u'\u2022', u''))
                index += 1
            except Exception as e:
                # 如果遇到错误,就将关键词从"used_keyword"列表中取出,重新获取此关键词进行拼接
                self.used_keyword.remove(keyword)
                print(e)

        # 重置已使用的关键词
        # self.used_keyword = []


if __name__ == '__main__':
    while True:
        file = FileSplitAndSplicing("")
        file.run()
        input_str = input("press Y(or y) + Enter to continue, press N(or n) + Enter to exit.")
        if input_str == "N" or input_str == "n":
            break
