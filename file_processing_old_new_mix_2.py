# coding:utf-8
import os
import re
import copy
import traceback
from docx import Document
from win32com.client import Dispatch
import codecs
import shutil
import openpyxl
from openpyxl import load_workbook
import random
from utils import Util
import docx
# import sys
# reload(sys)
# sys.setdefaultencoding('utf-8')


class File_processing5():
    """
    新文章和旧文章的拼接
    要求每篇文章中新段落和旧段落的比例2:1
    """

    def __init__(self, path):
        self.util = Util()
        self.document = docx
        # 域名
        self.read_path = path
        self.domain_name = "_".join(path.split("/")[-1].split("_")[:3])

        self.start_keyword = 0  # 关键词开始的位置
        self.end_keyword = int(path.split("/")[-1].split("_")[-1])  # 关键词结束的位置
        # self.special_keyword = '苏州'
        self.used_keyword = []

        # 保存路径
        self.save_article_path = r'./data/save_path/{}_articles'.format(self.domain_name)
        self.save_img_path = r'./data/save_path/{}_imgs'.format(self.domain_name)
        ####################   打包   ##########################
        # self.save_article_path = r'../data/save_path/{}_articles'.format(self.domain_name)
        # self.save_img_path = r'../data/save_path/{}_imgs'.format(self.domain_name)
        ########################################################
        # self.domain_name = 'uni_technology'
        # self.keywords_num = 180  # 关键词数量
        # 已使用的关键字
        self.used_keyword = []
        # 未使用的段落
        # self.unused_paragraphs = []
        # 已使用的图片
        # self.used_pictures = []
        # 所有的段落
        # self.paragraphs = []
        # self.keywords = self.read_xlsx(self.read_path + '\keyword.xlsx')
        self.used_articles = list()


    def get_keyword(self):
        """
        获取关键字(不能重复)
        :return:
        """
        try:
            unused_keyword = list(set(self.keywords) ^ set(self.used_keyword))
            if len(unused_keyword) == 0:
                return None
            keyword = random.choice(unused_keyword)
            self.used_keyword.append(keyword)
            return keyword
        except:
            traceback.print_exc()

    def operate_picture(self, filepath):
        """
        处理图片
        :param filepath:
        :return: 所有图片的路径
        """
        try:
            imgs = list()
            # for file in os.listdir(self.util.to_gbk(filepath)):
            #     img = file.decode('gbk').encode('utf-8')
            #     # imgs.append(os.path.join(filepath, img))
            #     imgs.append(img)
            for file in os.listdir(filepath):
                imgs.append(file)
            return imgs
        except:
            print('operate picture error', traceback.print_exc())

    def get_article_len(self, article):
        """
        求文章长度
        """
        article_len = 0
        for i in article:
            # article_len += len(i.decode('utf8'))
            article_len += len(i)
        return article_len

    def split_article(self, file_path, file):
        """
        将文章拆分为首段/中段/尾段
        """
        try:
            paragraphs = list()
            final_paragraphs = list()
            if file_path.endswith('txt'):
                # with open(file_path, 'r') as f:
                with open(file_path, 'r', encoding="UTF-8") as f:
                    paragraphs = f.readlines()
            elif file_path.endswith("docx"):
                # document = Document(u'{}'.format(file))
                document = Document(file_path)
                paragraphs = [p.text for p in document.paragraphs if p.text != '\n' and p != '\n' and p is not None and len(p.text) > 1]
            else:  # 读取doc
                word = Dispatch('Word.Application')  # 打开word应用程序
                # word = DispatchEx('Word.Application') # 启动独立的进程
                word.Visible = 0  # 后台运行,不显示
                word.DisplayAlerts = 0  # 不警告
                dir_path = os.path.dirname(os.path.abspath(file)) + "\\" + file_path
                doc = word.Documents.Open(dir_path)
                # doc.SaveAs(os.path.splitext(dir_path)[0] + '.docx', 12, False, "", True, "", False, False, False, False)
                for para in doc.paragraphs:
                    paragraphs.append(para.Range.Text)
                doc.Close()

            # 处理文档中的非正常回车
            for p in paragraphs:
                ps = p.split("\n\t")
                for pp in ps:
                    final_paragraphs.append(pp)

            start = final_paragraphs[0]
            middle = final_paragraphs[1: -1]
            end = final_paragraphs[-1]
            return start, middle, end
        except:
            traceback.print_exc()
            print("split_article error")
            doc.Close()

    def get_all_paragraphs(self, floder):
        """
        获取文章所有段落
        """
        try:
            start_paragraph_list = list()  # 存放所有首段段落
            middle_paragraph_list = list()  # 存放所有中段段落
            end_paragraph_list = list()  # 存放所有尾段段落
            filepath = self.read_path + '/' + floder
            file_list = [file for file in self.util.get_file_list(filepath) if file]

            for file in file_list:
                t_filepath = filepath + '/' + file
                start_paragraph, middle_paragraph, end_paragraph = self.split_article(t_filepath, file)
                if not start_paragraph or not middle_paragraph or not end_paragraph:
                    continue
                start_paragraph_list.append(start_paragraph)
                middle_paragraph_list.extend(middle_paragraph)
                end_paragraph_list.append(end_paragraph)
            return start_paragraph_list, middle_paragraph_list, end_paragraph_list
        except:
            print("get_all_paragraphs error", traceback.print_exc())

    # def get_keywords(self):
    #     """
    #     获取关键词
    #     举例:现需要150个关键词,含有"苏州"的关键词优先,
    #     如果含有"苏州"的关键词超过150个,则取前150个,
    #     如果含有"苏州"的关键词不到150,则取完这些词还要再取一些普通关键词凑够150个
    #     """
    #     all_keywords = self.read_xlsx()  # 所有的关键词
    #     special_keywords = [kw for kw in all_keywords if self.special_keyword in kw]  # 特殊关键词,如:含有"苏州"的关键词
    #     all_keywords = [kw for kw in all_keywords if self.special_keyword not in kw]  # 不含有"苏州"的所有的关键词
    #
    #     keywords_num = self.end_keyword - self.start_keyword  # 需要的关键词个数
    #     if len(special_keywords) >= keywords_num:
    #         needed_keywords = special_keywords[0:keywords_num]
    #         remaining_keywords = special_keywords[keywords_num:]
    #         remaining_keywords.extend(all_keywords)
    #     else:
    #         needed_keywords = copy.deepcopy(special_keywords)
    #         needed_keywords.extend(all_keywords[0: keywords_num-len(special_keywords)])
    #         remaining_keywords = all_keywords[(keywords_num-len(special_keywords)):]
    #     return needed_keywords, remaining_keywords

    # def write_unused_keywords_csv(self, keywords):
    #     """
    #     将未使用的关键词写入csv
    #     :return:
    #     """
    #     with open(self.save_article_path+'/unuserd_keywords.csv', 'w')as f:
    #         for kw in keywords:
    #             f.write(kw + '\n')

    def article_4_1(self, old_start_ps, new_middle_ps, old_end_ps):
        """
        1.old_start + 4 * new_middle + old_end
        """
        start_paragraph = random.sample(old_start_ps, 1)
        try:
            middle_paragraphs = random.sample(new_middle_ps, 4)
        except:
            try:
                middle_paragraphs = random.sample(new_middle_ps, 3)
            except:
                middle_paragraphs = random.sample(new_middle_ps, 2)
        end_paragraph = random.sample(old_end_ps, 1)
        article_ps = copy.deepcopy(start_paragraph)  # 文章
        article_ps.extend(middle_paragraphs)
        article_ps.extend(end_paragraph)
        return article_ps

    def article_4_2(self, old_start_ps, new_middle_ps, old_middle_ps, new_end_ps):
        """
        2.old_start + 3 * new_middle + 1 * old_middle + new_end
        """
        start_paragraph = random.sample(old_start_ps, 1)
        middle_paragraphs = random.sample(new_middle_ps, 3)
        middle_paragraphs.extend(random.sample(old_middle_ps, 1))
        end_paragraph = random.sample(new_end_ps, 1)
        article_ps = copy.deepcopy(start_paragraph)  # 文章
        article_ps.extend(middle_paragraphs)
        article_ps.extend(end_paragraph)
        return article_ps

    def article_4_3(self, new_start_ps, new_middle_ps, old_middle_ps, old_end_ps):
        """3.new_start + 3 * new_middle + 1 * old_middle + old_end"""
        start_paragraph = random.sample(new_start_ps, 1)
        middle_paragraphs = random.sample(new_middle_ps, 3)
        middle_paragraphs.extend(random.sample(old_middle_ps, 1))
        end_paragraph = random.sample(old_end_ps, 1)
        article_ps = copy.deepcopy(start_paragraph)  # 文章
        article_ps.extend(middle_paragraphs)
        article_ps.extend(end_paragraph)
        return article_ps

    def article_4_4(self, new_start_ps, new_middle_ps, old_middle_ps, new_end_ps):
        """4.new_start + 2 * new_middle + 2 * old_middle + new_end"""
        start_paragraph = random.sample(new_start_ps, 1)
        middle_paragraphs = random.sample(new_middle_ps, 2)
        middle_paragraphs.extend(random.sample(old_middle_ps, 2))
        end_paragraph = random.sample(new_end_ps, 1)
        article_ps = copy.deepcopy(start_paragraph)  # 文章
        article_ps.extend(middle_paragraphs)
        article_ps.extend(end_paragraph)
        return article_ps

    def random_article(self, old_start, old_middle, old_end, new_start, new_middle, new_end):
        """
        任意使用一种方法拼接文章
        """
        try:
            methods = [self.article_4_1(old_start, new_middle, old_end),
                       self.article_4_2(old_start, new_middle, old_middle, old_end),
                       self.article_4_3(new_start, new_middle, old_middle, new_end),
                       self.article_4_4(new_start, new_middle, old_middle, new_end)]
            return random.choice(methods)
        except:
            traceback.print_exc()

    def run(self):
        """
        拼接，生成一篇文章
        """
        if not os.path.exists(self.save_article_path):
            os.mkdir(self.save_article_path)
        if not os.path.exists(self.save_img_path):
            os.mkdir(self.save_img_path)
        # self.keywords, remaining_keywords = self.get_keywords()  # 获取关键词

        # 将剩余的关键词写入表格
        # self.write_unused_keywords_csv(remaining_keywords)

        img_list = list()
        old_start_paragraph_list = list()
        old_all_mid_list = list()
        old_end_paragraph_list = list()
        new_start_paragraph_list = list()
        new_all_mid_list = list()
        new_end_paragraph_list = list()

        file_dir_list = self.util.get_file_dir(self.read_path)  # 获取所有文件夹
        for floder in file_dir_list:  # 生成所有段落, 获取所有图片
            floder_path = self.read_path + '/' + floder
            # 旧文章段落
            if floder == 'old' and os.path.isdir(floder_path):
                old_start_paragraph_list, old_all_mid_list, old_end_paragraph_list = self.get_all_paragraphs(floder)
                old_all_mid_list = old_all_mid_list if len(old_all_mid_list) < 50 else random.sample(old_all_mid_list, 30)
            # 新文章段落
            elif floder == 'new' and os.path.isdir(floder_path):
                new_start_paragraph_list, new_all_mid_list, new_end_paragraph_list = self.get_all_paragraphs(floder)
                new_all_mid_list = new_all_mid_list if len(new_all_mid_list) < 50 else random.sample(new_all_mid_list, 30)
            # 图片
            elif floder == 'img':
                t_filepath = self.read_path + '/' + floder
                img_list = self.operate_picture(t_filepath)  # 获取所有图片
                for img in img_list:
                    shutil.copy(u"{}".format(self.read_path + '/' + floder + '/' +img), self.save_img_path)
            elif "xlsx" in floder:
                self.keywords = self.util.read_xlsx(floder_path)[self.start_keyword: self.end_keyword]  # 获取关键词

        # 拼接文章
        index = 1
        while True:
            keyword = self.util.get_keyword(self.keywords, self.used_keyword)  # 每一篇文章使用一个关键词
            if keyword is None:  # 关键词使用完之后退出循环
                break
            print(index, keyword)

            # 随机抽取文章，要求文章字数在730~870
            # while True:
            #     article = random.choice(article_list)  # 随机抽一篇文章
            #     article_len = self.get_article_len(article)
            #     if 730 < article_len < 870:
            #         break

            try:
                while True:  # 确保不会出现重复文章
                    article = self.random_article(old_start_paragraph_list, old_all_mid_list, old_end_paragraph_list,
                                                  new_start_paragraph_list, new_all_mid_list, new_end_paragraph_list)
                    if article not in self.used_articles:
                        self.used_articles.append(article)
                        break
                temp_article = copy.deepcopy(article)  # 深拷贝，对新数据进行处理，不改变原数据
                img = random.sample(img_list, 2)  # 随机取两张图
                article_str = ''
                ####  段落 -- 对每一段进行处理
                for num in range(len(temp_article)):
                    if num == 0 or num == len(temp_article) - 1:  # 添加首段/尾段
                        temp_article[num] = self.util.insert_keyword(keyword, temp_article[num])  # 插入关键词
                        article_str += '<p>%s</p>\n' % temp_article[num]
                    elif num == 1:  # 添加第二段，并插入一张图片
                        article_str += '<p>%s</p>\n' % temp_article[num]
                        article_str += '<p><img src="{imgpath}/%s_imgs/%s"></p>\n' % (
                            self.domain_name, img[0])  # 注意修改站点名称
                    elif num == 3:  # 添加第四段，并插入一张图片
                        article_str += '<p>%s</p>\n' % temp_article[num]
                        article_str += '<p><img src="{imgpath}/%s_imgs/%s"></p>\n' % (
                            self.domain_name, img[1])  # 注意修改站点名称
                    else:  # 添加第三段
                        article_str += '<p>%s</p>\n' % temp_article[num]
                save_path = self.save_article_path + '/' + '{}.txt'.format(keyword)
                try:
                    self.util.write_article(save_path, article_str)
                except:
                    self.util.write_article(save_path, article_str.replace(u'\u200b', u'').replace(u'\xa0', u'').replace(u'\u2022', u''))
                index += 1
            except Exception as e:
                # 如果遇到错误,就将关键词从"used_keyword"列表中取出,这样就可以重新获取此关键词进行拼接
                self.used_keyword.remove(keyword)
                print(e)


if __name__ == "__main__":
    file_processing = File_processing5("./data/read_path/china_aiying_mix_1_150")
    file_path = "./data/read_path/china_aiying_mix_1_150/new/otc搬运机器人是近代主动控制范畴呈现的一项高新技能.docx"
    file = "otc搬运机器人是近代主动控制范畴呈现的一项高新技能.docx"
    start, middle, end = file_processing.split_article(file_path, file)
